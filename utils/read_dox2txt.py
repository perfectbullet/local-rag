from docx import Document



def readDocx(file_path):
    doc = Document(file_path)
    content = ''
    # 每一段的编号、内容
    for i in range(len(doc.paragraphs)):
        istr = str(i) + " " + doc.paragraphs[i].text + "\n"
        content += istr
    # 表格
    tbs = doc.tables
    for tb in tbs:
        # 行
        for row in tb.rows:
            # 列
            for cell in row.cells:
                cell_str = cell.text + "\t"
                content += cell_str
            content += '\n'
    return content


if __name__ == '__main__':
    content = readDocx('../test_files/医疗器械临床试验机构自查报告.docx')
    print(content)
